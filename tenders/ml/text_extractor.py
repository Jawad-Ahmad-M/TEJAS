"""
Text Extractor Module
=====================
Extract text from PDF, Word, and text files for tender processing.
"""

import os
import re
from io import BytesIO


def extract_text_from_file(file):
    """
    Extract text from uploaded file (PDF, Word, or text).
    
    Args:
        file: Django UploadedFile object
        
    Returns:
        str: Extracted text content
    """
    filename = file.name.lower()
    file_content = file.read()
    
    # Reset file pointer
    file.seek(0)
    
    if filename.endswith('.pdf'):
        return _extract_from_pdf(file_content)
    elif filename.endswith('.docx'):
        return _extract_from_docx(file_content)
    elif filename.endswith('.doc'):
        return _extract_from_doc(file_content)
    elif filename.endswith('.txt'):
        return file_content.decode('utf-8', errors='ignore')
    else:
        raise ValueError(f"Unsupported file type: {filename}")


def _extract_from_pdf(content):
    """Extract text from PDF bytes."""
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(BytesIO(content))
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return "\n".join(text_parts)
    except ImportError:
        raise ImportError("PyPDF2 is required for PDF extraction. Install with: pip install PyPDF2")
    except Exception as e:
        raise ValueError(f"Error extracting PDF text: {str(e)}")


def _extract_from_docx(content):
    """Extract text from DOCX bytes."""
    try:
        from docx import Document
        
        doc = Document(BytesIO(content))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    except ImportError:
        raise ImportError("python-docx is required for Word extraction. Install with: pip install python-docx")
    except Exception as e:
        raise ValueError(f"Error extracting DOCX text: {str(e)}")


def _extract_from_doc(content):
    """Extract text from DOC bytes (legacy Word format)."""
    # For .doc files, we try to use textract or return error
    try:
        import textract
        return textract.process(BytesIO(content)).decode('utf-8', errors='ignore')
    except ImportError:
        raise ImportError("For .doc files, please convert to .docx format or install textract")
    except Exception as e:
        raise ValueError(f"Error extracting DOC text: {str(e)}")


def parse_text_to_columns(text):
    """
<<<<<<< HEAD
    Parse extracted text into the required column format.
    
    This is a basic parser - you may need to customize based on your document structure.
    
    Args:
        text: Raw extracted text
        
    Returns:
        dict: Dictionary with keys matching required columns
    """
    # Initialize with empty values
=======
    Parse extracted text into the required column format using robust regex.
    """
>>>>>>> bd1274c (Added Chat and rafactored code)
    columns = {
        'Title': '',
        'Authority': '',
        'Object_Description': '',
        'CPV': '',
        'Estimated_Value': '',
        'Award_Criteria': '',
        'Conditions': ''
    }
    
<<<<<<< HEAD
    # Try to extract sections based on common patterns
    text_lower = text.lower()
    
    # Title - usually first significant line or after "title:" pattern
    title_match = re.search(r'(?:title|subject|tender name)[:\s]*([^\n]+)', text, re.IGNORECASE)
    if title_match:
        columns['Title'] = title_match.group(1).strip()
    else:
        # Take first non-empty line as title
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        if lines:
            columns['Title'] = lines[0][:200]  # Limit title length
    
    # Authority
    auth_match = re.search(r'(?:contracting authority|authority|organization|awarding body)[:\s]*([^\n]+)', text, re.IGNORECASE)
    if auth_match:
        columns['Authority'] = auth_match.group(1).strip()
    
    # CPV code (8-digit number)
    cpv_match = re.search(r'(?:cpv|common procurement vocabulary)[:\s]*(\d{8})', text, re.IGNORECASE)
    if not cpv_match:
        cpv_match = re.search(r'\b(\d{8})\b', text)  # Any 8-digit number
    if cpv_match:
        columns['CPV'] = cpv_match.group(1)
    
    # Estimated Value
    value_patterns = [
        r'(?:estimated value|contract value|budget|value)[:\s]*[€$£]?\s*([\d,\.]+)',
        r'[€$£]\s*([\d,\.]+)',
        r'([\d,\.]+)\s*(?:EUR|USD|GBP)',
    ]
    for pattern in value_patterns:
        value_match = re.search(pattern, text, re.IGNORECASE)
        if value_match:
            columns['Estimated_Value'] = value_match.group(1).replace(',', '')
            break
    
    # Object Description - main body text
    desc_match = re.search(r'(?:description|object|scope)[:\s]*([^\n]+(?:\n[^\n]+)*)', text, re.IGNORECASE)
    if desc_match:
        columns['Object_Description'] = desc_match.group(1).strip()[:5000]
    else:
        # Use large portion of text as description
        columns['Object_Description'] = text[:5000]
    
    # Award Criteria
    criteria_match = re.search(r'(?:award criteria|evaluation criteria|criteria)[:\s]*([^\n]+(?:\n[^\n]+)*)', text, re.IGNORECASE)
    if criteria_match:
        columns['Award_Criteria'] = criteria_match.group(1).strip()[:2000]
    
    # Conditions
    cond_match = re.search(r'(?:conditions|requirements|eligibility)[:\s]*([^\n]+(?:\n[^\n]+)*)', text, re.IGNORECASE)
    if cond_match:
        columns['Conditions'] = cond_match.group(1).strip()[:2000]
    
=======
    # 1. Title Extraction
    title_match = re.search(r'(?:title|tender name|subject|project)[:\s]+([^\n]+)', text, re.IGNORECASE)
    if title_match:
        columns['Title'] = title_match.group(1).strip()
    else:
        # Fallback: Find the first non-trivial line
        for line in text.split('\n'):
            line = line.strip()
            if len(line) > 10 and not any(k in line.lower() for k in ['cpv', 'value', 'criteria']):
                columns['Title'] = line[:200]
                break
    
    # 2. Authority / Organization
    auth_match = re.search(r'(?:contracting authority|authority|organization|awarding body|issued by)[:\s]+([^\n]+)', text, re.IGNORECASE)
    if auth_match:
        columns['Authority'] = auth_match.group(1).strip()

    # 3. CPV Code (8 digits, potentially hyphenated or with checksum)
    cpv_match = re.search(r'(?:cpv|procurement vocabulary codes?)[:\s]+(\d{8}(?:-\d)?)', text, re.IGNORECASE)
    if not cpv_match:
        cpv_match = re.search(r'\b(\d{8})\b', text)
    if cpv_match:
        columns['CPV'] = cpv_match.group(1).split('-')[0]

    # 4. Estimated Value (Currencies: €, $, £, or EUR, USD, GBP)
    val_match = re.search(r'(?:estimated value|budget|total value|contract amount)[:\s]+(?:[€$£]|eur|usd|gbp)?\s*([\d\s,\.]+)', text, re.IGNORECASE)
    if val_match:
        # Clean numeric string: remove spaces and handle decimal points correctly
        val_str = re.sub(r'[\s,]', '', val_match.group(1).strip())
        columns['Estimated_Value'] = val_str

    # 5. Section Extraction (Description, Criteria, Conditions)
    sections = {
        'Object_Description': [r'(?:object description|scope of work|description)[:\s]+', r'(?:award criteria|award)'],
        'Award_Criteria': [r'(?:award criteria|evaluation criteria|criteria)[:\s]+', r'(?:conditions|eligibility)'],
        'Conditions': [r'(?:conditions|eligibility requirements|participation conditions)[:\s]+', r'(?:procedure|time limit|deadline)']
    }

    for key, markers in sections.items():
        start_pattern = markers[0]
        end_pattern = markers[1]
        
        # Combine patterns and use re.IGNORECASE flag instead of (?i) prefix
        # This avoids "global flags not at the start" error
        match = re.search(f"{start_pattern}(.*?)(?={end_pattern}|$)", text, re.DOTALL | re.IGNORECASE)
        if match:
            columns[key] = match.group(1).strip()[:5000 if key == 'Object_Description' else 2000]

    # Post-process cleanup
    for key in columns:
        if isinstance(columns[key], str):
            # Remove multiple spaces/newlines
            columns[key] = re.sub(r'\s+', ' ', columns[key]).strip()

>>>>>>> bd1274c (Added Chat and rafactored code)
    return columns
